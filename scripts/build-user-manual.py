from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "manual"
DOCX = OUTPUT / "Manual-do-Usuario-CaixaSimples-Bratec.docx"
BLUE = "2563EB"
INK = "0F172A"
MUTED = "475569"
LIGHT = "E8EEF5"


def font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def fixed_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)


def mark_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    font(run, 9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    font(r, 22, True, INK)
    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(16)
        r = s.add_run(subtitle)
        font(r, 11, color=MUTED)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        font(r, 11, True)
        r = p.add_run(text[len(bold_prefix):])
        font(r)
    else:
        font(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        font(p.add_run(item))


def callout(doc, title, text, fill="EFF6FF"):
    table = doc.add_table(rows=1, cols=1)
    fixed_table(table, [9360])
    mark_header(table.rows[0])
    shade(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(title + "\n"), 11, True, BLUE)
    font(p.add_run(text), 10.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def chapter(doc, title, kicker):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(kicker.upper()), 9, True, BLUE)
    add_title(doc, title)


def add_picture(doc, path, alt, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    doc_pr = run._element.xpath(".//wp:docPr")[0]
    doc_pr.set("descr", alt)
    p.paragraph_format.space_after = Pt(5)


def build():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font(header.add_run("CAIXA NO CONTROLE  |  MANUAL DO USUÁRIO"), 8.5, True, MUTED)
    page_number(section.footer.paragraphs[0])

    # Editorial cover
    doc.add_paragraph().paragraph_format.space_after = Pt(36)
    logo = ROOT / "src-tauri" / "icons" / "icon.png"
    add_picture(doc, logo, "Ícone azul do aplicativo CaixaSimples - Bratec", 2.15)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    font(p.add_run("MANUAL DO USUÁRIO"), 10, True, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run("CaixaSimples - Bratec"), 30, True, INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Gestão financeira simples para pequenos negócios"), 14, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Versão 1.2.0-beta.1  •  Agosto de 2026  •  BratecInfo"), 10, True, BLUE)

    chapter(doc, "Antes de começar", "1 · visão geral")
    add_body(doc, "O CaixaSimples - Bratec organiza caixa, contas, vendas, metas, relatórios, licenças e backups em um aplicativo desktop. O funcionamento principal é offline: seus dados financeiros ficam no computador.")
    callout(doc, "Responsabilidade compartilhada", "O aplicativo ajuda na gestão, mas não substitui orientação contábil, fiscal ou jurídica. Mantenha backups em outro dispositivo.")
    doc.add_heading("Requisitos", level=2)
    add_bullets(doc, ["Windows 10 ou 11, 64 bits.", "Permissão para instalar aplicativos no perfil do usuário.", "Espaço adicional para o WebView2 offline e para os backups."])
    doc.add_heading("Mapa do aplicativo", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    fixed_table(table, [2700, 6660])
    mark_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, ("Área", "Finalidade")):
        shade(cell, LIGHT); font(cell.paragraphs[0].add_run(value), 10, True)
    for area, purpose in [("Visão geral", "Indicadores e alertas"), ("Movimentações", "Receitas, despesas e ajustes"), ("Vendas", "Orçamentos confirmados e parcelas"), ("Relatórios", "Análises e exportações"), ("Backup", "Continuidade e diagnóstico")]:
        cells = table.add_row().cells
        for cell, value in zip(cells, (area, purpose)): font(cell.paragraphs[0].add_run(value), 10)

    chapter(doc, "Instalação no Windows", "2 · distribuição")
    add_bullets(doc, ["Execute CaixaSimples - Bratec_1.2.0-beta.1_x64-setup.exe.", "Leia e aceite o contrato de licença.", "Confirme a pasta sugerida e o atalho opcional da área de trabalho.", "Abra o aplicativo pelo menu Iniciar > BratecInfo."])
    callout(doc, "Instalação sem internet", "O instalador inclui o WebView2 Runtime offline. Em versões atuais do Windows ele normalmente já está presente.")
    doc.add_heading("Aviso de assinatura", level=2)
    add_body(doc, "Enquanto a BratecInfo não configurar um certificado comercial, o Windows poderá mostrar “Editor desconhecido”. Baixe o instalador somente de um canal autorizado e compare o SHA-256 com o arquivo SHA256SUMS.txt do release.")
    doc.add_heading("Atualizar ou desinstalar", level=2)
    add_body(doc, "Instalar uma versão mais nova sobre a anterior preserva o banco. Ao desinstalar, escolha Não na pergunta separada sobre dados para mantê-los. Escolha Sim apenas depois de criar um backup e quando quiser apagar tudo permanentemente.")

    chapter(doc, "Primeiro acesso", "3 · configuração")
    add_body(doc, "O assistente prepara a empresa, preferências e o administrador local em etapas curtas. Campos obrigatórios são validados antes de avançar.")
    add_bullets(doc, ["Informe nome e tipo do negócio.", "Escolha caixa ou competência como visão padrão.", "Cadastre a conta, o saldo e a data de abertura.", "Selecione categorias e formas de pagamento.", "Defina uma meta mensal opcional.", "Crie uma senha local com pelo menos 12 caracteres."])
    callout(doc, "Senha local", "A senha é transformada com Argon2id e nunca é exibida ou registrada em logs.")
    doc.add_heading("Avaliação e demonstração", level=2)
    add_body(doc, "A avaliação permite 50 operações financeiras, sem expiração por data. Vendas parceladas, parcelamentos e transferências contam uma vez por ação. Após o limite, consultas, relatórios, exportações e backups permanecem disponíveis. Os dados demonstrativos são uma escolha independente e podem ser removidos sem apagar os registros reais.")

    chapter(doc, "Conheça a interface", "4 · navegação")
    screenshot = ROOT / "docs" / "phase-9-audit" / "06-dashboard-final-completo.png"
    if screenshot.exists():
        add_picture(doc, screenshot, "Tela Visão geral com indicadores financeiros, alertas e navegação lateral")
    add_body(doc, "O menu lateral organiza as rotinas por assunto. Em telas menores ele pode ser recolhido. A faixa superior informa a área atual e o aviso de avaliação aparece quando aplicável.")
    add_bullets(doc, ["Use Tab e Shift+Tab para navegar sem mouse.", "Pressione Enter ou Espaço para ativar botões e links.", "Use Esc para fechar diálogos.", "Mensagens de sucesso, alerta e erro informam o resultado de cada ação."])

    chapter(doc, "Cadastros e dados demonstrativos", "5 · base operacional")
    add_body(doc, "Cadastros bem preenchidos reduzem retrabalho nas vendas, contas e relatórios.")
    doc.add_heading("Clientes e fornecedores", level=2)
    add_body(doc, "Registre nome, papéis, documento, contatos, endereço, observações e etiquetas. O sistema sugere possíveis duplicidades antes de salvar.")
    doc.add_heading("Produtos, serviços e referências", level=2)
    add_body(doc, "Informe código, descrição, preços em reais e unidade. Em Cadastros básicos, mantenha categorias, contas e formas de pagamento ativas.")
    doc.add_heading("Pacote demonstrativo", level=2)
    add_body(doc, "Em Configurações, carregue exemplos de cliente, fornecedor, itens, movimentações e venda parcelada. Use Remover dados demonstrativos quando terminar; somente registros marcados como demonstração serão excluídos.")

    chapter(doc, "Receitas, despesas e obrigações", "6 · financeiro")
    add_body(doc, "Use Movimentações para receitas, despesas, aportes, retiradas, ajustes e recorrências. Datas de emissão, competência, vencimento e liquidação têm funções diferentes.")
    doc.add_heading("Liquidação", level=2)
    add_body(doc, "Em Contas a receber ou Contas a pagar, selecione uma obrigação e informe valor, data, conta e forma de pagamento. Liquidações parciais mantêm o saldo restante em aberto.")
    add_bullets(doc, ["Desconto reduz o valor líquido.", "Juros, multa e taxa aumentam o valor líquido.", "Cancelamento e estorno exigem motivo e mantêm auditoria.", "Recorrências geram ocorrências controladas sem duplicidade."])
    callout(doc, "Valores confiáveis", "A aplicação armazena moeda como inteiro em centavos. Transferências não entram como receita ou despesa.")

    chapter(doc, "Vendas e parcelas", "7 · comercial")
    add_body(doc, "Cadastre cliente e itens, crie a venda e escolha recebimento imediato, futuro, parcelado ou misto. Antes de confirmar, revise totais, descontos, taxas e datas.")
    add_bullets(doc, ["Recebimento imediato liquida a parcela na conta escolhida.", "Venda futura cria uma conta a receber.", "Parcelamento distribui o valor e preserva o total em centavos.", "Recibo em PDF apresenta dados da empresa, cliente, itens e totais."])
    callout(doc, "Cancelamento seguro", "Se houver recebimentos, siga a orientação apresentada para estornar ou regularizar antes de cancelar.", "FEF2F2")

    chapter(doc, "Caixa, competência, metas e relatórios", "8 · gestão")
    add_body(doc, "A Visão geral apresenta saldo, entradas, saídas, resultado, vencidos e próximos compromissos. O Fluxo de caixa detalha saldos por conta e projeções.")
    doc.add_heading("Dois regimes", level=2)
    add_bullets(doc, ["Caixa: considera quando o dinheiro foi recebido ou pago.", "Competência: considera quando a receita ou despesa pertence ao negócio."])
    doc.add_heading("Metas e relatórios", level=2)
    add_body(doc, "Defina metas mensais de receita, limite de despesas, resultado, vendas e novos clientes. Nos relatórios, escolha período, regime e filtros; revise a prévia antes de exportar PDF ou CSV.")

    chapter(doc, "Backup, restauração e atualização", "9 · continuidade")
    add_body(doc, "Configure uma pasta de backup fora do computador sempre que possível. A proteção por senha usa derivação Argon2id e criptografia autenticada; a senha não pode ser recuperada.")
    add_bullets(doc, ["Crie backups antes de atualizações e mudanças importantes.", "Confira histórico, tamanho e checksum.", "Na restauração, valide empresa, versão e proteção antes de confirmar.", "A restauração cria uma cópia preventiva e reinicia o aplicativo."])
    callout(doc, "Regra 3-2-1", "Mantenha três cópias, em dois tipos de mídia, com uma cópia fora do computador.")
    doc.add_heading("Atualização assinada", level=2)
    add_body(doc, "Pacotes .cncupd são verificados por assinatura, checksum, versão mínima do banco e compatibilidade da licença antes de abrir o instalador.")

    chapter(doc, "Assinatura e segurança", "10 · confiança")
    add_body(doc, "Em Configurações > Meu plano, escolha Mensal (R$ 9,90 por mês) ou Anual (R$ 99,90 por ano). O pagamento abre no ambiente seguro do Mercado Pago; o aplicativo não solicita nem armazena dados do cartão. Depois do pagamento, volte e selecione Verificar pagamento.")
    add_body(doc, "Uma autorização temporária assinada permite o uso offline por até 7 dias no plano mensal ou 30 dias no anual, sempre limitada ao período pago. Quando houver internet, use Verificar pagamento para renovar. Arquivos .cnclic existem apenas para clientes com licença legada já emitida.")
    add_bullets(doc, ["Não compartilhe o arquivo de licença fora dos dispositivos autorizados.", "Não envie backups ou diagnósticos sem revisar o destino.", "Não edite o banco SQLite manualmente.", "Use apenas instaladores e atualizações com hashes publicados."])
    callout(doc, "Privacidade", "O pacote de diagnóstico contém versão, sistema e logs técnicos, mas exclui banco financeiro, senhas e chaves privadas.")

    chapter(doc, "Solução de problemas", "capítulo 11 · suporte")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    fixed_table(table, [3300, 6060])
    mark_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, ("Situação", "O que fazer")):
        shade(cell, LIGHT); font(cell.paragraphs[0].add_run(value), 10, True)
    rows = [
        ("Aplicativo não abre", "Reinicie o Windows e confira se o antivírus isolou o executável."),
        ("Ação recusada", "Revise campos obrigatórios, datas, valores e a mensagem apresentada."),
        ("Alerta no banco", "Não edite o arquivo; gere diagnóstico e restaure um backup validado."),
        ("Backup não restaura", "Confirme senha, extensão .cncbak, integridade e versão compatível."),
        ("Assinatura não ativa", "Confira a internet e use Verificar pagamento em Meu plano. Se o backend estiver temporariamente indisponível, o lease válido continua funcionando."),
    ]
    for situation, action in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, (situation, action)): font(cell.paragraphs[0].add_run(value), 10)
    doc.add_heading("Ao pedir suporte", level=2)
    add_body(doc, "Informe a versão 1.2.0-beta.1, a edição do Windows, os passos para reproduzir e a mensagem exata. Anexe o pacote de diagnóstico somente se desejar.")

    chapter(doc, "Referência rápida", "12 · consulta")
    doc.add_heading("Atalhos", level=2)
    add_bullets(doc, ["Tab / Shift+Tab — próximo ou controle anterior", "Enter / Espaço — ativar", "Esc — fechar diálogo", "Ctrl+F no leitor de PDF — pesquisar neste manual"])
    doc.add_heading("Dados locais", level=2)
    add_body(doc, "Banco interno preservado: %APPDATA%\\br.com.bratecinfo.caixanocontrole\\caixa-no-controle.db")
    add_body(doc, "Logs: %LOCALAPPDATA%\\br.com.bratecinfo.caixanocontrole\\logs")
    callout(doc, "Antes de apagar", "Crie um backup protegido, copie-o para outro dispositivo e confirme o checksum.", "FFF7ED")
    doc.add_heading("Sobre", level=2)
    add_body(doc, "CaixaSimples - Bratec 1.2.0-beta.1 · BratecInfo · Manual revisado em 10 de agosto de 2026.")

    # Keep page furniture consistent in all generated sections.
    for sec in doc.sections:
        sec.header.is_linked_to_previous = True
        sec.footer.is_linked_to_previous = True

    doc.core_properties.title = "Manual do usuário — CaixaSimples - Bratec"
    doc.core_properties.subject = "Guia de instalação e operação do CaixaSimples - Bratec 1.2.0-beta.1"
    doc.core_properties.author = "BratecInfo"
    doc.core_properties.keywords = "financeiro, caixa, vendas, backup, offline"
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
